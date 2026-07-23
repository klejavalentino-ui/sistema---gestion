import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_zecat_guide():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Header / Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("GUÍA DE INTEGRACIÓN ZECAT WEB & DATAMARGEN ERP")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("Conexión de tienda joymomerch.productosconlogo.com, Catálogo Mayorista, Precios y Stock por Ubicación Web")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.add_paragraph()

    # Introduction box
    table_intro = doc.add_table(rows=1, cols=1)
    table_intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_intro = table_intro.cell(0, 0)
    set_cell_background(cell_intro, "F8FAFC")
    set_cell_margins(cell_intro, top=140, bottom=140, left=180, right=180)
    
    p_intro = cell_intro.paragraphs[0]
    r_intro_bold = p_intro.add_run("📌 Objetivo de la Integración:\n")
    r_intro_bold.bold = True
    r_intro_bold.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p_intro.add_run(
        "Esta guía técnica explica paso a paso cómo conectar la plataforma web de merchandising de Zecat "
        "(joymomerch.productosconlogo.com) con Datamargen ERP para Jomo Indumentaria (jomoindumentaria@gmail.com). "
        "Permite sincronizar automáticamente el catálogo de artículos, códigos SKU, costos mayoristas, márgenes de venta "
        "y stock disponible bajo la sucursal/ubicación 'Web' en la sección de Inventario."
    )

    doc.add_paragraph()

    # Section 1: Credenciales necesarias
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Datos y Credenciales requeridas de Zecat")
    r_h1.font.size = Pt(14)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_desc = doc.add_paragraph()
    p_desc.add_run("Para completar la vinculación en Datamargen, necesitás solicitar a Zecat dos claves de identificación únicas:")

    # Table of credentials
    table_cred = doc.add_table(rows=3, cols=2)
    table_cred.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = table_cred.rows[0].cells
    set_cell_background(hdr[0], "0F172A")
    set_cell_background(hdr[1], "0F172A")
    r0 = hdr[0].paragraphs[0].add_run("Credencial / Dato")
    r0.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1 = hdr[1].paragraphs[0].add_run("¿Qué es y cómo obtenerlo?")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    row1 = table_cred.rows[1].cells
    set_cell_background(row1[0], "F1F5F9")
    row1[0].paragraphs[0].add_run("Token API 2.0 Partner").bold = True
    p1 = row1[1].paragraphs[0]
    p1.add_run(
        "Es el Token de seguridad (Bearer Key) que autoriza al sistema a consultar la API 2.0 de Zecat.\n"
        "• Cómo pedirlo: Enviar un email a tu ejecutivo de cuenta o a sistemas de Zecat solicitando el 'Token API de Partner para integración con ERP'."
    )

    row2 = table_cred.rows[2].cells
    set_cell_background(row2[0], "FFFFFF")
    row2[0].paragraphs[0].add_run("Código / ID de Partner Zecat").bold = True
    p2 = row2[1].paragraphs[0]
    p2.add_run(
        "Es tu número de cuenta único de distribuidor/revendedor oficial en la plataforma de Zecat.\n"
        "• Cómo encontrarlo: Figura en tu panel de control de Zecat Partner (extranet), en el encabezado de tus facturas de compra mayorista a Zecat o en el contrato de tu tienda joymomerch.productosconlogo.com."
    )

    for row in table_cred.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

    doc.add_paragraph()

    # Section 2: Paso a paso de configuración en Datamargen
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. Paso a Paso de Carga en Datamargen ERP")
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    steps = [
        ("Paso 1: Ingresar a la sección Zecat Web", 
         "En el menú lateral de Datamargen, hacer clic en la solapa '🏬 Zecat Web'."),
        ("Paso 2: Completar el Formulario de Conexión", 
         "1. URL de la Tienda Web: Verificar que esté cargada https://joymomerch.productosconlogo.com\n"
         "2. Token API Partner Zecat: Pegar el token de seguridad provisto por Zecat.\n"
         "3. Código / ID de Partner Zecat: Ingresar el ID de cuenta de revendedor Zecat."),
        ("Paso 3: Guardar Credenciales", 
         "Hacer clic en el botón 'Guardar Credenciales y Conectar'. El sistema cambiará el badge a '🟢 Conectado con Zecat API'."),
        ("Paso 4: Ejecutar la Sincronización del Catálogo", 
         "Hacer clic en 'Sincronizar Catálogo'. Datamargen descargará las familias de productos (mochilas, indumentaria, botellas, lapiceras), sus precios mayoristas de costo y el stock disponible.")
    ]

    for step_title, step_desc in steps:
        p_step = doc.add_paragraph()
        r_st = p_step.add_run(step_title + "\n")
        r_st.bold = True
        r_st.font.size = Pt(11.5)
        r_st.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        p_step.add_run(step_desc)
        p_step.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    # Section 3: Integración con Inventario y Ubicación WEB
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. Gestión de Inventario por Ubicación (Depósito Casa vs. Web)")
    r_h3.font.size = Pt(14)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_inv = doc.add_paragraph()
    p_inv.add_run(
        "Para la cuenta de Jomo Indumentaria (jomoindumentaria@gmail.com), el sistema organiza el inventario en 2 ubicaciones distintas:\n\n"
        "• 📦 Ubicación 'Web': Recibe de forma automática todo el stock de los productos del catálogo de Zecat sincronizados de la tienda online.\n"
        "• 🏠 Ubicación 'Depósito Casa': Permite registrar el stock físico propio que guardás en tu casa (muestras, compras directas o stock para entrega inmediata).\n\n"
        "Al cotizar en la sección de 'Presupuestos' o al facturar una venta, vas a poder elegir de qué ubicación descontar el stock según la preferencia del cliente."
    )

    doc.add_paragraph()

    # Footer note
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ft = p_foot.add_run("Documento generado por Datamargen ERP • www.datamargen.com")
    r_ft.font.size = Pt(9)
    r_ft.font.italic = True
    r_ft.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(r"c:\Users\kleja\OneDrive\Escritorio\Mazo\Guia_Paso_a_Paso_Integracion_Zecat_Datamargen.docx")
    print("Documento Word creado con éxito.")

if __name__ == "__main__":
    create_zecat_guide()
